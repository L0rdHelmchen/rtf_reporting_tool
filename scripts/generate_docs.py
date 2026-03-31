#!/usr/bin/env python3
"""
RTF Reporting Tool – Projektdokumentation Generator
Erzeugt eine Word-Dokumentation (.docx) für das Projekt.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # Bundesbank-Blau
    return h

def add_info_table(doc, rows):
    """Zweispaltige Key/Value-Tabelle."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(row.cells[0], 'EFF6FF')
        row.cells[1].text = value
    doc.add_paragraph()

def add_code_block(doc, code, lang=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6F)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_tech_table(doc, headers, data):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        hdr_row.cells[i].text = h
        hdr_row.cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr_row.cells[i], '1A56DB')
        hdr_row.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_data in data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    doc.add_paragraph()

# ── Document ───────────────────────────────────────────────────────────────────

doc = Document()

# Seitenränder
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Titelseite ─────────────────────────────────────────────────────────────────

title = doc.add_heading('RTF Meldungs-Tool', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

sub = doc.add_paragraph('Projektdokumentation – Technische Architektur & Implementierung')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(13)
sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Version: 1.0  |  Stand: {datetime.date.today().strftime('%d.%m.%Y')}  |  "
             f"Taxonomie: RTF 2023-12\n")
meta.add_run("Deutsche Bundesbank – XBRL-basiertes Meldewesen-Tool")
meta.runs[-1].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ── Inhaltsverzeichnis (manuell) ───────────────────────────────────────────────

add_heading(doc, 'Inhaltsverzeichnis', 1)
toc_entries = [
    ('1', 'Projektübersicht'),
    ('2', 'Technologie-Stack'),
    ('3', 'Systemarchitektur'),
    ('4', 'Projektstruktur'),
    ('5', 'Backend – Kernkomponenten'),
    ('6', 'Frontend – Aufbau & Navigation'),
    ('7', 'XBRL-Taxonomie-Integration'),
    ('8', 'Validierungsarchitektur'),
    ('9', 'Datenbank-Schema'),
    ('10', 'Docker & Deployment'),
    ('11', 'API-Referenz'),
    ('12', 'Bekannte Bugs & Fixes'),
    ('13', 'Offene Punkte / Roadmap'),
]
for num, entry in toc_entries:
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(f"{num}   {entry}")
    run.font.size = Pt(11)

doc.add_page_break()

# ── 1. Projektübersicht ────────────────────────────────────────────────────────

add_heading(doc, '1  Projektübersicht', 1)

doc.add_paragraph(
    'Das RTF Meldungs-Tool ist eine webbasierte Plattform zur Erfassung, Validierung und '
    'Einreichung von aufsichtsrechtlichen Meldungen nach dem XBRL-Standard der Deutschen '
    'Bundesbank (RTF-Taxonomie 2023-12). Es richtet sich an Institute, die gemäß KWG zur '
    'Meldung von Risikotragfähigkeit, Gruppenangaben, Liquiditätsbewertung und Kapitalplanung '
    'verpflichtet sind.'
)

add_heading(doc, '1.1  Ziele', 2)
goals = [
    'Digitale Erfassung aller RTF-Meldeformulare (41 Formulare, 84 Konzepte)',
    'Automatische XBRL-Validierung gegen die offizielle Bundesbank-Taxonomie',
    'Offline-Auswertung von 2.374 Value Assertions (Quersummen, Plausibilitäten)',
    'Prüfung von 9 Existence Assertions (Pflichtfelder DBL/STG)',
    'Generierung gültiger XBRL-Instanzdokumente für die Einreichung',
    'Mandantenfähigkeit mit Row-Level-Security pro Institut',
    'Vollständige Audit-Trail-Erfassung aller Formularänderungen',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

add_heading(doc, '1.2  Projekteckdaten', 2)
add_info_table(doc, [
    ('Taxonomie',           'Deutsche Bundesbank RTF 2023-12'),
    ('Taxonomie-Pfad',      'RTF_Validierungsdateien/www.bundesbank.de/…/rtf-2023-12/2023-12-31/'),
    ('Formulare',           '41 RTF-Formulare aus tab/tab.xsd + tab-lab-de.xml'),
    ('Metrik-Konzepte',     '84 Konzepte aus dict/met/met.xsd'),
    ('Enum-Domains',        '52 Domains vorgeladen (FX=183, IC=251, WZ=107, …)'),
    ('Value Assertions',    '2.374 Regeln (99,2 % auswertbar) aus val/vr-v*.xml'),
    ('Existence Assertions','9 Regeln (DBL: 8, STG: 1) aus val/vr-e*.xml'),
    ('Login (Dev)',         'admin / admin123'),
    ('DB (Dev)',            'rtf_user / rtf_password @ localhost:5432 / rtf_reporting'),
])

# ── 2. Technologie-Stack ───────────────────────────────────────────────────────

add_heading(doc, '2  Technologie-Stack', 1)

add_tech_table(doc,
    ['Schicht', 'Technologie', 'Version', 'Zweck'],
    [
        ['Frontend',  'React',                    '18',       'SPA-Framework'],
        ['Frontend',  'Vite',                     '5',        'Build-Tool & Dev-Server'],
        ['Frontend',  'MUI (Material UI)',         '5',        'UI-Komponentenbibliothek'],
        ['Frontend',  'Redux Toolkit',             'aktuell',  'State Management'],
        ['Frontend',  'react-hook-form + Zod',     'aktuell',  'Formular-Validierung'],
        ['Frontend',  'i18next',                   'aktuell',  'Internationaliserung'],
        ['Frontend',  'Luxon + MUI X DatePicker',  'aktuell',  'Datumseingabe (di5)'],
        ['Backend',   'Fastify',                   'v5',       'HTTP-Framework'],
        ['Backend',   'TypeScript + tsx watch',    'aktuell',  'Sprache / Dev-Reload'],
        ['Backend',   'libxmljs2',                 'aktuell',  'XBRL-XML-Parsing'],
        ['Datenbank', 'PostgreSQL',                '15',       'Persistenz + RLS'],
        ['Cache',     'Redis',                     '7',        'Session / Rate-Limit'],
        ['Shared',    '@rtf-tool/shared',          'intern',   'Typen/Konstanten (nur TS)'],
        ['DevOps',    'Docker Compose',            'v2',       'lokale Entwicklungsumgebung'],
    ]
)

# ── 3. Systemarchitektur ───────────────────────────────────────────────────────

add_heading(doc, '3  Systemarchitektur', 1)

doc.add_paragraph(
    'Das System ist als Monorepo mit drei Paketen organisiert: frontend, backend und shared. '
    'Alle Dienste laufen in Docker-Containern mit benannten Volumes für node_modules.'
)

add_heading(doc, '3.1  Container-Übersicht', 2)
add_tech_table(doc,
    ['Container', 'Name', 'Port', 'Status'],
    [
        ['Frontend (Vite HMR)', 'rtf-frontend', '5173', 'Healthy'],
        ['Backend (Fastify)',   'rtf-backend',  '3001', 'Healthy'],
        ['PostgreSQL',          'rtf-postgres', '5432', 'Healthy'],
        ['Redis',               'rtf-redis',    '6379', 'Healthy'],
    ]
)

add_heading(doc, '3.2  Kommunikationsfluss', 2)
doc.add_paragraph(
    'Browser → Vite-Proxy (/api/v1/*) → Fastify-Backend → PostgreSQL / Redis\n'
    'Alle API-Requests werden vom Vite-Dev-Server als Proxy weitergeleitet. '
    'Im Production-Betrieb übernimmt nginx diese Aufgabe.'
)

add_heading(doc, '3.3  Authentifizierung', 2)
doc.add_paragraph(
    'JWT-basierte Authentifizierung. Tokens werden im localStorage gespeichert '
    'und bei jedem Request als Bearer-Token mitgesendet. '
    'Rate-Limit: 1.000 Requests / 15 min (erhöht von Default 100, da alle '
    'Vite-Proxy-Requests dieselbe Docker-IP teilen).'
)

# ── 4. Projektstruktur ─────────────────────────────────────────────────────────

add_heading(doc, '4  Projektstruktur', 1)

add_code_block(doc,
"""rtf_reporting_tool/
├── backend/
│   ├── src/
│   │   ├── data/                  # Generierte JSON-Daten
│   │   │   ├── required-fields.json   # Existence Assertions (9 Regeln)
│   │   │   └── value-rules.json       # Value Assertions (2.374 Regeln)
│   │   ├── lib/
│   │   │   ├── database.ts        # PostgreSQL Pool + RLS-Helfer
│   │   │   └── logger.ts
│   │   ├── routes/
│   │   │   ├── forms.ts           # /api/v1/forms
│   │   │   ├── auth.ts            # /api/v1/auth
│   │   │   └── reportingPeriods.ts
│   │   ├── services/
│   │   │   ├── XBRLSchemaParser.ts       # Taxonomie-Parser (Hauptklasse)
│   │   │   ├── XBRLGeneratorService.ts   # XBRL-Instanz-Generator
│   │   │   ├── XBRLValidationService.ts  # Validierungs-Orchestrator
│   │   │   └── ValueAssertionEvaluator.ts # Value-Assertion-Engine
│   │   └── server.ts
│   └── package.json
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── forms/             # FormRenderer, FormFieldComponents
│   │   │   └── layout/            # Sidebar, Header
│   │   ├── pages/
│   │   │   ├── forms/             # FormsPage, FormEditorPage
│   │   │   └── dashboard/
│   │   ├── services/              # formsApi.ts, api.ts
│   │   └── store/slices/          # Redux-Slices
│   └── vite.config.ts
├── shared/src/
│   ├── types.ts
│   └── index.ts                   # FORM_CATEGORY_NAMES_DE etc.
├── scripts/
│   ├── parse-existence-assertions.mjs  # Generiert required-fields.json
│   └── parse-value-assertions.mjs      # Generiert value-rules.json
├── database/init/                 # SQL-Migrationsskripte
└── RTF_Validierungsdateien/       # Bundesbank XBRL-Taxonomie (gemountet)""")

# ── 5. Backend – Kernkomponenten ───────────────────────────────────────────────

add_heading(doc, '5  Backend – Kernkomponenten', 1)

add_heading(doc, '5.1  XBRLSchemaParser', 2)
doc.add_paragraph(
    'Zentrale Klasse zur Verarbeitung der RTF-Taxonomie. Beim Start wird '
    'parseRTFTaxonomy() einmalig aufgerufen; das Ergebnis wird gecacht.'
)
add_tech_table(doc,
    ['Methode', 'Beschreibung'],
    [
        ['parseRTFTaxonomy()',         'Lädt tab.xsd, met.xsd, alle dom/*.xsd und required-fields.json'],
        ['getForms()',                 'Gibt Map<id, FormDefinition> zurück (41 Formulare)'],
        ['getFormDefinition(id)',      'Einzelnes Formular mit Sections und Fields'],
        ['validateFormData(id, data)', 'Prüft Pflichtfelder (Schema + Existence Assertions)'],
        ['buildRequiredFieldsSection()', 'Erstellt Pflichtangaben-Sektion für DBL/STG'],
        ['loadRequiredFields()',       'Lädt required-fields.json → requiredFields Map'],
    ]
)

add_heading(doc, '5.2  ValueAssertionEvaluator', 2)
doc.add_paragraph(
    'Wertet die 2.374 Value Assertions aus der RTF-Taxonomie gegen Formulardaten aus. '
    'Wird lazy beim ersten validateFormData()-Aufruf geladen (value-rules.json, 1,6 MB).'
)
add_tech_table(doc,
    ['Pattern', 'Anzahl', 'Beispiel', 'Abdeckung'],
    [
        ['conditional', '1.623', 'if (exists($a)) then ($b = true())', '✓ implementiert'],
        ['simple',      '678',   '$a >= 0, $a <= $b, $a = $b',         '✓ implementiert'],
        ['iaf',         '68',    'iaf:numeric-equal($a, iaf:sum(…))',  '✓ implementiert'],
        ['regex',       '5',     "matches($a, '^\\d{8}$')",            '✓ implementiert'],
        ['unknown',     '0',     '–',                                   '–'],
    ]
)
doc.add_paragraph(
    'Gesamt: 2.356 von 2.374 Regeln vollständig auswertbar (99,2 %). '
    '18 Iterator-Regeln (keine expliziten Zellkoordinaten) werden übersprungen.'
)

add_heading(doc, '5.3  XBRLValidationService – Validierungsreihenfolge', 2)
steps = [
    '1. Schema-Validierung (Pflichtfelder, Datentypen) via XBRLSchemaParser.validateFormData()',
    '2. Business Rules (CAP_001, CAP_002, … – formularspezifische Regeln)',
    '3. XBRL Value Assertions via ValueAssertionEvaluator.evaluate(formCode, formData)',
    '4. Cross-Form-Konsistenz (formularübergreifende Prüfungen)',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '5.4  Datenbankzugriff', 2)
doc.add_paragraph(
    'Alle Schreibzugriffe laufen über db.transactionAs(userId, callback). '
    'Dabei wird PostgreSQL Row-Level-Security via set_config(\'app.current_user_id\', userId) aktiviert. '
    'Nutzer können ausschließlich Daten ihrer eigenen Institution lesen und schreiben.'
)

# ── 6. Frontend ────────────────────────────────────────────────────────────────

add_heading(doc, '6  Frontend – Aufbau & Navigation', 1)

add_heading(doc, '6.1  Seitenstruktur', 2)
add_tech_table(doc,
    ['Route', 'Komponente', 'Funktion'],
    [
        ['/dashboard',                    'DashboardPage',   'Übersicht, Statistiken'],
        ['/forms',                        'FormsPage',       'Formularliste mit Kategorie-/Statusfilter'],
        ['/forms?category=GRP',           'FormsPage',       'Gefilterte Ansicht: Gruppenangaben'],
        ['/forms/:id/edit/:period',       'FormEditorPage',  'Formular bearbeiten & speichern'],
        ['/forms/:id',                    'FormDetailPage',  'Formular-Detailansicht'],
        ['/xbrl',                         'XBRLPage',        'XBRL-Export'],
        ['/institution',                  'InstitutionPage', 'Institutionsverwaltung'],
    ]
)

add_heading(doc, '6.2  Formular-Rendering', 2)
doc.add_paragraph(
    'FormRenderer.tsx rendert Sections und Fields dynamisch aus der FormDefinition. '
    'Jedes Feld wird anhand seines dataType auf eine spezialisierte Komponente gemappt:'
)
add_tech_table(doc,
    ['dataType', 'Komponente', 'Beschreibung'],
    [
        ['si6',  'TextField',       'Freitext / String'],
        ['mi1',  'CurrencyField',   'Geldbeträge (mit Tausendertrennzeichen)'],
        ['pi2',  'PercentageField', 'Prozentwerte'],
        ['ii3',  'IntegerField',    'Ganzzahlen'],
        ['di5',  'DateFieldDI5',    'Datum (Luxon DateTime → ISO-String)'],
        ['bi7',  'BooleanField',    'Checkbox (true/false)'],
        ['ei8',  'EnumField',       'Auswahlliste (XBRL-Codelisten)'],
    ]
)

add_heading(doc, '6.3  Kategorie-Filter', 2)
doc.add_paragraph(
    'Sidebar-Links navigieren zu /forms?category=GRP etc. FormsPage liest den '
    'URL-Parameter via useSearchParams() und sendet ihn als Query-Parameter an die API. '
    'Das Backend filtert die Formularliste server-seitig. '
    'Die Badge-Zahlen in der Sidebar werden beim ersten Laden dynamisch von GET /forms abgerufen.'
)

# ── 7. XBRL-Taxonomie ─────────────────────────────────────────────────────────

add_heading(doc, '7  XBRL-Taxonomie-Integration', 1)

add_heading(doc, '7.1  Verzeichnisstruktur (Taxonomie)', 2)
add_code_block(doc,
"""RTF_Validierungsdateien/www.bundesbank.de/…/rtf-2023-12/2023-12-31/
├── tax.xsd                  # Entry Point
├── tab/
│   ├── tab.xsd              # Formulardefinitionen (41 tgXXX-Elemente)
│   └── tab-lab-de.xml       # Deutsche Labels
├── dict/
│   ├── met/met.xsd          # 84 Metrik-Konzepte
│   └── dom/{CODE}/
│       ├── mem.xsd          # Enum-Werte
│       └── mem-lab-de.xml   # Deutsche Labels
└── val/
    ├── vr-e*.xml            # Existence Assertions (9 Dateien)
    ├── vr-e*-err-de.xml     # Fehlermeldungen (DE)
    ├── vr-v*.xml            # Value Assertions (2.374 Dateien)
    └── vr-v*-err-de.xml     # Fehlermeldungen (DE)""")

add_heading(doc, '7.2  Generierte Datendateien', 2)
doc.add_paragraph(
    'Die Parser-Skripte in scripts/ lesen die Taxonomie-Dateien einmalig ein und erzeugen '
    'optimierte JSON-Dateien für das Backend. Diese müssen nach Taxonomie-Aktualisierungen '
    'neu generiert werden:'
)
add_code_block(doc,
"""# Existence Assertions (Pflichtfelder DBL + STG)
node scripts/parse-existence-assertions.mjs
# → backend/src/data/required-fields.json

# Value Assertions (Quersummen, Plausibilitäten)
node scripts/parse-value-assertions.mjs
# → backend/src/data/value-rules.json""")

# ── 8. Validierungsarchitektur ─────────────────────────────────────────────────

add_heading(doc, '8  Validierungsarchitektur', 1)

add_heading(doc, '8.1  Existence Assertions (Pflichtfelder)', 2)
doc.add_paragraph(
    'Nur 9 Existence-Assertion-Dateien in der gesamten RTF-Taxonomie. '
    'Sie definieren Pflichtfelder für DBL (8 Felder) und STG (1 Feld). '
    'Alle anderen Formulare haben keine Existence Assertions.'
)
add_tech_table(doc,
    ['Template', 'Felder', 'Beispiel'],
    [
        ['DBL', '8', 'x010 (Institutsname), x020 (Kreditgeber-ID), x040 (Stichtag), …'],
        ['STG', '1', 'x010_x010 (Frequenz der Berichterstattung)'],
    ]
)

add_heading(doc, '8.2  Value Assertions (Arithmetik & Plausibilität)', 2)
add_tech_table(doc,
    ['Template', 'Regeln', 'Template', 'Regeln'],
    [
        ['ILAAP',   '229', 'RSK',    '459'],
        ['KPL',     '312', 'RDP-R',  '584'],
        ['RDP-BH',  '278', 'RDP-BI', '257'],
        ['RDP-BW',  '113', 'STG',    '7'],
        ['GRP',     '19',  'STA',    '24'],
        ['CRS',     '25',  'STKK',   '57'],
        ['DBL',     '2',   'RTFK',   '8'],
    ]
)

add_heading(doc, '8.3  Zellschlüssel-Logik', 2)
doc.add_paragraph(
    'XBRL-Fakten werden durch Dimensionswerte adressiert. '
    'Der Evaluator berechnet den Zellschlüssel für jede Variable:'
)
add_code_block(doc,
"""effectiveRow = variable.row ?? globalRow
effectiveCol = variable.col ?? globalCol

Wenn row + col:  Schlüssel = "x050_x030"
Wenn nur row:    Schlüssel = "x050"
Wenn nur col:    Schlüssel = "x030"
Wenn keiner:     Iterator (wird übersprungen)""")

# ── 9. Datenbank-Schema ────────────────────────────────────────────────────────

add_heading(doc, '9  Datenbank-Schema', 1)

add_heading(doc, '9.1  Kern-Tabellen', 2)
add_tech_table(doc,
    ['Tabelle', 'Zweck', 'Wichtige Spalten'],
    [
        ['institutions',        'Meldepflichtige Institute',         'id, name, bafin_id'],
        ['users',               'Benutzer pro Institut',             'id, institution_id, role, email'],
        ['form_instances',      'Gespeicherte Formulare',            'id, form_code, reporting_period, status, form_data (JSONB)'],
        ['form_instance_history','Audit-Trail aller Änderungen',     'form_instance_id, changed_by, change_type, previous_data, new_data'],
        ['reporting_periods',   'Meldezeiträume (jährlich)',         'id, year, start_date, end_date, is_active'],
    ]
)

add_heading(doc, '9.2  form_instances – Status-Workflow', 2)
add_code_block(doc,
"""draft → in_review → submitted → accepted
                  ↘ rejected""")

add_heading(doc, '9.3  Row-Level-Security', 2)
doc.add_paragraph(
    'Alle Tabellen haben RLS-Policies, die auf institution_id filtern. '
    'Der App-Benutzer (rtf_app) hat kein SUPERUSER-Recht. '
    'Der Nutzerkontext wird per set_config(\'app.current_user_id\', userId, true) '
    'am Anfang jeder Transaktion gesetzt.'
)

# ── 10. Docker & Deployment ────────────────────────────────────────────────────

add_heading(doc, '10  Docker & Deployment', 1)

add_heading(doc, '10.1  Entwicklungsumgebung', 2)
add_code_block(doc,
"""# Start (alle Container)
docker compose up -d

# Rebuild nach Code-Änderungen
docker compose up --build -d backend frontend

# Logs
docker compose logs -f backend

# Datenbank-Shell
docker compose exec postgres psql -U rtf_user -d rtf_reporting""")

add_heading(doc, '10.2  Wichtige Umgebungsvariablen', 2)
add_tech_table(doc,
    ['Variable', 'Wert (Dev)', 'Beschreibung'],
    [
        ['DATABASE_URL',          'postgresql://rtf_user:rtf_password@postgres:5432/rtf_reporting', 'DB-Verbindung'],
        ['REDIS_URL',             'redis://redis:6379',    'Redis-Verbindung'],
        ['JWT_SECRET',            '(in .env)',              'Token-Signatur'],
        ['RATE_LIMIT_MAX_REQUESTS','1000',                  'Erhöht (Docker-IP-Sharing)'],
        ['NODE_ENV',              'development',            'Umgebung'],
    ]
)

add_heading(doc, '10.3  XBRL-Dateien', 2)
doc.add_paragraph(
    'Die Taxonomie-Dateien liegen im Host-Verzeichnis RTF_Validierungsdateien/ '
    'und werden als Read-Only-Volume in den Backend-Container gemountet. '
    'Im Container sind sie unter /app/RTF_Validierungsdateien/ verfügbar.'
)

# ── 11. API-Referenz ───────────────────────────────────────────────────────────

add_heading(doc, '11  API-Referenz', 1)

add_heading(doc, '11.1  Authentifizierung', 2)
add_tech_table(doc,
    ['Methode', 'Endpunkt', 'Beschreibung'],
    [
        ['POST', '/api/v1/auth/login',   'Login → JWT-Token'],
        ['GET',  '/api/v1/auth/me',      'Aktueller Nutzer'],
        ['POST', '/api/v1/auth/logout',  'Logout'],
        ['POST', '/api/v1/auth/refresh', 'Token erneuern'],
    ]
)

add_heading(doc, '11.2  Formulare', 2)
add_tech_table(doc,
    ['Methode', 'Endpunkt', 'Beschreibung'],
    [
        ['GET',  '/api/v1/forms',                          'Alle Formulare (mit ?category=, ?search=)'],
        ['GET',  '/api/v1/forms/:id',                      'Formular-Definition'],
        ['GET',  '/api/v1/forms/:id/instance',             'Gespeicherte Instanz'],
        ['POST', '/api/v1/forms/:id/instance',             'Neue Instanz erstellen'],
        ['PUT',  '/api/v1/forms/:id/instance',             'Instanz speichern (+ Validierung)'],
        ['POST', '/api/v1/forms/:id/submit',               'Formular einreichen'],
        ['POST', '/api/v1/forms/:id/validate',             'Nur validieren (kein Speichern)'],
        ['GET',  '/api/v1/forms/:id/export/xbrl',         'XBRL-Export'],
        ['GET',  '/api/v1/reporting-periods',             'Meldezeiträume'],
    ]
)

# ── 12. Bekannte Bugs & Fixes ──────────────────────────────────────────────────

add_heading(doc, '12  Bekannte Bugs & Fixes (Projektverlauf)', 1)

add_tech_table(doc,
    ['#', 'Problem', 'Ursache', 'Fix'],
    [
        ['1',  'this.xxx = fn crash',          'this=undefined in strict-mode async',       'Lokale const statt this-Zuweisung'],
        ['2',  'RTF_BASE_PATH falsch',          '4× ../ statt 3×',                          'Pfad korrigiert'],
        ['3',  'schemaParser.onReady crash',    'Fataler Fehler im Hook',                   'console.warn statt throw'],
        ['4',  'reply.getResponseTime() fehlt', 'Fastify v5 API geändert',                  'reply.elapsedTime'],
        ['5',  'getForms() fehlt',              'Falsche Methode',                          '[...schemaParser.getForms().values()]'],
        ['6',  'Healthcheck 404',               'Falsche URL',                              '/health/ping'],
        ['7',  'shared-Paket Syntaxfehler',     'Vite kein CJS→ESM für @rtf-tool/shared',  'optimizeDeps.include in vite.config.ts'],
        ['8',  'date-fns nicht auflösbar',      'Redundanter AdapterDateFns',               'Entfernt – AdapterLuxon in main.tsx reicht'],
        ['9',  'PersistGate hängt',             'persistReducer fehlte',                    'combineReducers + persistReducer'],
        ['10', 'HTTP 429 Rate-Limit',           'Docker-IP-Sharing',                        'RATE_LIMIT_MAX_REQUESTS=1000'],
        ['11', 'Infinite GET /auth/me',         'authLoading in useEffect-Deps',            'Aus Dep-Array entfernt'],
        ['12', 'forms.length crash',            'Falsches Data-Unwrapping',                 'inner = data?.data ?? data'],
        ['13', 'XBRL XPath Namespace-Error',    'libxmljs2 erwartet Plain Object',          'get ns() Getter + find(…, this.ns)'],
        ['14', '0 Formulare',                   'Einzel-Form-XSDs fehlen im Download',      'parseFormsFromTabXsd() Fallback'],
        ['15', '/reporting-periods 404',        'Route fehlte',                             'reportingPeriods.ts erstellt'],
        ['16', 'Datum: received object',        'DatePicker liefert Luxon-Objekt, Zod z.date()', 'z.string() + DateTime.toISODate()'],
        ['17', 'Kategorie-Filter greift nicht', 'getFormCategory() englische Strings; GET / filtert nicht', 'Kurzcodes + Server-seitiger Filter'],
        ['18', 'Badge-Zahlen hardcodiert',      'Statische Werte in Sidebar',               'Dynamic via getFormCounts() API-Call'],
    ]
)

# ── 13. Roadmap ────────────────────────────────────────────────────────────────

add_heading(doc, '13  Offene Punkte / Roadmap', 1)

add_heading(doc, '13.1  Kurzfristig (Pflicht)', 2)
short = [
    'Server-seitige Paginierung für GET /forms (aktuell: alle Formulare auf einmal)',
    'Fehlerbehandlung im FormEditorPage bei fehlendem Berichtszeitraum',
    'E-Mail-Validierung im Registrierungsformular',
    'Produktions-Build testen (Dockerfile.prod + docker-compose.prod.yml)',
]
for s in short:
    doc.add_paragraph(s, style='List Bullet')

add_heading(doc, '13.2  Mittelfristig', 2)
mid = [
    'Iterator-Value-Assertions implementieren (18 übersprungene Regeln)',
    'XBRL-Instanzvalidierung gegen Bundesbank-Einreichungsregeln',
    'Dashboard mit echten Statistiken aus der Datenbank',
    'Rollenbasierte Freigabe-Workflow (4-Augen-Prinzip)',
    'Mehrsprachigkeit (EN) über i18next ausbauen',
]
for m in mid:
    doc.add_paragraph(m, style='List Bullet')

add_heading(doc, '13.3  Langfristig', 2)
long = [
    'Automatischer Taxonomie-Update-Prozess (RTF 2024+)',
    'Direkteinreichung an Bundesbank-API',
    'Vergleich von Meldungen über mehrere Perioden',
    'Mobile-optimiertes Layout',
]
for l in long:
    doc.add_paragraph(l, style='List Bullet')

# ── Footer / Speichern ─────────────────────────────────────────────────────────

doc.add_page_break()
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final.add_run(
    f"RTF Meldungs-Tool – Projektdokumentation v1.0\n"
    f"Stand: {datetime.date.today().strftime('%d.%m.%Y')} | "
    f"Deutsche Bundesbank Taxonomie 2023-12"
)
final.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
final.runs[0].font.size = Pt(9)

# Speichern
out_path = os.path.join(
    os.path.dirname(__file__), '..', 'RTF_Meldungs_Tool_Projektdokumentation.docx'
)
doc.save(out_path)
print(f"Gespeichert: {os.path.abspath(out_path)}")
